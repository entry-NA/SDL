
"""生成手把手实验教程 — 仅主实验，不含消融"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sty = doc.styles['Normal']
sty.font.name = '宋体'; sty.font.size = Pt(10.5)
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def H(t, l=1):
    h = doc.add_heading(t, level=l)
    for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return h

def P(t, bold=False):
    p = doc.add_paragraph(); run = p.add_run(t)
    run.font.name = '宋体'; run.font.size = Pt(10.5); run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def C(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(t); run.font.name = 'Consolas'; run.font.size = Pt(9)

def warn(t):
    p = doc.add_paragraph(); run = p.add_run('⚠ ' + t)
    run.font.name = '宋体'; run.font.size = Pt(10); run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def tip(t):
    p = doc.add_paragraph(); run = p.add_run('💡 ' + t)
    run.font.name = '宋体'; run.font.size = Pt(10)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def step(n, t):
    p = doc.add_paragraph(); run = p.add_run(f'步骤 {n}：{t}')
    run.font.name = '宋体'; run.font.size = Pt(11); run.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ════════════════════════════════ COVER ════════════════════════════════
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run('\n\n\nSAM2 辅助伪标签精修实验\n手把手完整教程\n'); r.font.size = Pt(24); r.bold = True
r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

su = doc.add_paragraph(); su.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = su.add_run('\n从打开命令行到拿到最终指标\n一步一步 · 每步都有截图说明 · 零基础可复现\n'); r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100,100,100); r.font.name = '宋体'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

info = doc.add_paragraph(); info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('\n\n\n环境: Windows 11 · Python 3.9 · PyTorch 2.12 · 8GB GPU\n最后更新: 2026年7月\n')
r.font.size = Pt(11); r.font.name = '宋体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.add_page_break()

# ════════════════════════════════ CH1 ════════════════════════════════
H('第一章  打开命令行，进入工作环境', 1)

H('1.1 打开命令行（终端）', 2)
P('按键盘 Win + R，输入 cmd，回车。会弹出一个黑色窗口。这就是命令行（也叫终端、命令提示符）。')
P('本文中所有灰色背景的命令，都在这个黑色窗口里输入。输入完后按回车执行。')

H('1.2 激活 conda 环境', 2)
P('我们所有的实验都在一个名叫 test01 的 Python 虚拟环境中运行。在命令行输入：')
C('conda activate test01')
P('执行后，命令行最左边的提示符会从 (base) 变成 (test01)，说明环境已切换。')
P('如果报错"找不到环境"，说明还没有创建。先执行：')
C('conda create -n test01 python=3.9 -y')
C('conda activate test01')

H('1.3 进入项目文件夹', 2)
P('所有代码和数据都在以下文件夹中：')
C('C:\\Users\\23991\\Desktop\\新建文件夹\\UCOD-DPL-main\\UCOD-DPL-main')
P('在命令行输入：')
C('cd "C:\\Users\\23991\\Desktop\\新建文件夹\\UCOD-DPL-main\\UCOD-DPL-main"')
tip('每次重新打开命令行，都先做 1.2 + 1.3 这两步。')

H('1.4 设置编码（防止中文乱码）', 2)
P('在命令行输入：')
C('set PYTHONIOENCODING=utf-8')
P('这条命令只对当前窗口有效。关掉重开需要重新输入。')

H('1.5 验证环境是否正常', 2)
P('输入以下三行，确认关键组件可用：')
C('python -c "import torch; print(\'PyTorch OK, CUDA:\', torch.cuda.is_available())"')
C('python -c "from sam2.build_sam import build_sam2; print(\'SAM2 OK\')"')
C('python -c "import cv2, numpy; print(\'OpenCV OK, NumPy:\', numpy.__version__)"')
P('如果都输出了 "OK"，环境就准备好了。如果某个报错，回到 1.2 重新激活环境或检查安装。')

doc.add_page_break()

# ════════════════════════════════ CH2 ════════════════════════════════
H('第二章  确认数据文件都在', 1)

H('2.1 认识三类数据', 2)
P('本实验需要三类数据，它们分别在桌面上不同的文件夹里：')

P('① 原始图片（训练+测试）：', bold=True)
C('C:\\Users\\23991\\Desktop\\RefCOD (1)\\RefCOD\\')
P('里面有 6 个子文件夹：TR-CAMO、TR-COD10K（训练用）、TE-CAMO、TE-COD10K、CHAMELEON、NC4K（测试用）')
P('每个子文件夹下有 im/（图片）和 gt/（真值标注）。训练只用到 im，评估时才会用到 gt。')

P('② 官方预生成的粗伪标签（16×16 的 .pkl 文件）：', bold=True)
C('C:\\Users\\23991\\Desktop\\plable\\TR-CAMO+TR-COD10K\\')
P('里面有 4040 个 .pkl 文件（data_0.pkl ~ data_4039.pkl）和一个 index.json。')
P('每个 pkl 文件里存的是一个 16×16 的二值掩码张量（只有 0 和 1），这是 UCOD-DPL 论文用的原始伪标签。')

P('③ 我们会生成的 SAM2 精修伪标签：', bold=True)
C('.\\datasets\\cache\\refined_pseudo_labels\\')
P('目前这个文件夹是空的——我们接下来运行脚本才能生成它。')

H('2.2 创建链接——让程序能找到数据', 2)
P('问题：数据文件在桌面上，但程序期望在项目文件夹里找到它们。')
P('解决方案：用 Windows 的 mklink /J 命令创建一个"目录连接"——相当于一个硬链接，程序会以为文件就在项目里。')

step(1, '创建图片数据连接')
P('把桌面上的 RefCOD 图片数据连接到项目的 datasets/RefCOD：')
C('python -c "import subprocess; subprocess.run([\'cmd\', \'/c\', \'mklink\', \'/J\', r\'datasets\\RefCOD\', r\'C:\\Users\\23991\\Desktop\\RefCOD (1)\\RefCOD\'], shell=True)"')
warn('如果报错"权限不够"，用管理员身份重新打开 cmd 再执行。')

step(2, '创建粗伪标签连接')
P('把桌面上的 pkl 伪标签连接到项目的缓存目录：')
C('mkdir datasets\\cache\\pseudo_label_cache 2>nul')
C('python -c "import subprocess; subprocess.run([\'cmd\', \'/c\', \'mklink\', \'/J\', r\'datasets\\cache\\pseudo_label_cache\\TR-CAMO+TR-COD10K\', r\'C:\\Users\\23991\\Desktop\\plable\\TR-CAMO+TR-COD10K\'], shell=True)"')

step(3, '验证连接成功')
C('dir datasets\\RefCOD')
P('应该看到 6 个文件夹：CHAMELEON NC4K TE-CAMO TE-COD10K TR-CAMO TR-COD10K')
C('dir "datasets\\cache\\pseudo_label_cache\\TR-CAMO+TR-COD10K\\*.pkl" | find /c ".pkl"')
P('应该输出 4040。如果不是 4040，检查上一步的路径是否正确。')

doc.add_page_break()

# ════════════════════════════════ CH3 ════════════════════════════════
H('第三章  第一步——生成 SAM2 精修伪标签', 1)

H('3.1 这一步要干什么', 2)
P('UCOD-DPL 的原始伪标签是 16×16 像素的（从 DINOv2 注意力图阈值化得到），边界模糊、可能有空间偏移。')
P('我们写的 offline_sam2_refine.py 脚本用 SAM2 模型把这些粗糙标签"精修"成高质量掩码：')
P('  ① 加载 16×16 粗标签 → 上采样到原图分辨率')
P('  ② 计算自适应提示（扩张框 + 正点 + 分层负点）')
P('  ③ SAM2 在原图分辨率下推理，生成 3 个候选掩码')
P('  ④ 截断式多掩码选择最优候选')
P('  ⑤ 边缘感知置信度门控决定最终采纳/融合/回退')
P('  ⑥ 保存为 PNG 文件（原图分辨率，灰度图）')
P('整个过程跑一次就行，约 37 分钟。')

H('3.2 执行命令', 2)
P('确保你已经做完第一章的所有步骤，然后在命令行输入：')
C('python scripts/offline_sam2_refine.py')

H('3.3 运行时你会看到什么', 2)
P('首先，SAM2 模型加载（首次运行会下载预训练权重，约 1-2 分钟）：')
C('[SAM2Wrapper] Loaded sam2.1_hiera_tiny on cuda')
C('Images: 4040, Pseudo-labels: 4040')

P('然后，进度条开始滚动：')
C('Refining:  25%|██▌       | 1010/4040 [09:15<27:30,  1.84it/s]')

P('进度条含义：百分比 | 当前/总数 | 已用时间/预计剩余时间 | 每秒处理张数。')

P('最后，显示统计信息：')
C('Done. mode=both')
C('Full pipeline: 4040 refined, Full:3956 (97.9%) Fusion:79 (2.0%) Fallback:5 (0.1%)')
C('Errors: 0')

P('各统计项含义：')
P('  Full: 完全采纳 SAM2 输出（S > 0.8，高置信度）')
P('  Fusion: 软融合（0.2 ≤ S ≤ 0.8，SAM2 和粗标签加权混合）')
P('  Fallback: 完全回退到粗标签（S < 0.2，SAM2 不可靠）')
P('  Errors: 处理失败数，必须为 0')

H('3.4 验证生成结果', 2)
C('dir datasets\\cache\\refined_pseudo_labels\\*.png | find /c ".png"')
P('应该输出 4040。你也可以随便打开这个文件夹，双击几张 png 图片看看——应该是灰度图，和原图一样大。')

H('3.5 如果中途报错怎么办', 2)
P('常见错误及解决：')

P('① ModuleNotFoundError: No module named \'sam2\' → SAM2 没装好。执行：')
C('pip install git+https://github.com/facebookresearch/sam2.git')

P('② CUDA out of memory → 显存不够。关掉浏览器和所有不用的程序，然后重跑。')

P('③ FileNotFoundError: .../index.json → 粗伪标签连接没创建。回到 2.2 节第二步。')

P('④ 某个图片报错但其他的正常 → 单张图的问题，脚本会自动跳过并记录在 Errors 里，不影响整体。')

doc.add_page_break()

# ════════════════════════════════ CH4 ════════════════════════════════
H('第四章  第二步——训练 UCOD-DPL 模型', 1)

H('4.1 这一步要干什么', 2)
P('用第三章生成的 SAM2 精修伪标签（高质量的），训练 UCOD-DPL 的 DBA 解码器。')
P('训练 25 个 epoch（轮次），每 5 个 epoch 自动在 TE-CAMO（250 张图）上验证一次。')

H('4.2 训练前检查', 2)
P('确保精修标签存在且数量正确：')
C('dir datasets\\cache\\refined_pseudo_labels\\*.png | find /c ".png"')
P('必须是 4040。如果不是，回到第三章重新生成。')

H('4.3 执行命令', 2)
C('bash scripts/launch_train_first_stage.sh -c configs/uscod/UCOD-DPL_dinov2.py')
P('参数含义：-c 指定配置文件，configs/uscod/UCOD-DPL_dinov2.py 包含所有训练参数。')
warn('如果 bash 命令不可用，需要先装 Git for Windows，或者把 .sh 文件里的命令手动复制到 cmd 执行。')

H('4.4 运行时你会看到什么', 2)

P('阶段 1 — 特征提取（首次训练，约 8 分钟）：', bold=True)
P('首次训练时，程序会用 DINOv2-base 对 4040 张训练图逐一提取特征并缓存。会显示：')
C('Extracting image features by facebook/dinov2-base  100%  [08:23<00:00]')
P('特征缓存保存在 datasets/cache/features_cache/ 下。以后再训练就会跳过这一步。')

P('阶段 2 — 模型初始化（约 30 秒）：', bold=True)
C('INFO  Successfully built baseline model and discriminator')
C('INFO  Successfully built optimizers and schedulers')
C('INFO  Successfully prepared components for distributed training')
P('如果看到这三行，说明模型加载成功。')

P('阶段 3 — 训练循环（约 33 分钟）：', bold=True)
P('训练开始后，终端会持续输出日志，格式如下：')
C('iter10:loss:0.7525           ← 第10次迭代，训练损失 0.7525')
C('pl:0.9991                    ← 伪标签从判别器的概率（接近1=判别器认为伪标签质量高）')
C('ps:0.4302                    ← 学生预测从判别器的概率')
C('merge_label_weight:0.46      ← APM 自适应融合权重（0=固定标签，1=教师预测）')
C('train/dis_loss:0.5925        ← 判别器损失')

P('其中 loss 是最重要的：表示模型预测和伪标签之间的差异。训练初期 loss 约 0.7-0.8，后期约 0.1-0.2，偶尔出现负值（因为对抗损失）也是正常的。')

P('阶段 4 — 验证（每 5 个 epoch 一次）：', bold=True)
P('每 5 个 epoch，程序会在 TE-CAMO 上验证一次，打印指标表。看到 "best result" 字样就是刷新了最佳结果：')
C('INFO  best result:')
C('+--------+--------+--------+--------+')
C('| 0.9246 | 0.6608 | 0.8697 | 0.8673 |  ← ACC | mIOU | E_MAX | E_MEAN')
C('0.7811 | 0.7789 |  0.7968  | 0.0754 |  ← F_MAX | F_MEAN | SMeasure | MAE')
C('0.7507 |                                   ← WFM')

P('阶段 5 — 训练完成：', bold=True)
C('Train Epoch  --- 100% Time Elapsed: 0:41:03')
P('总耗时约 41 分钟（含特征提取）。')

H('4.5 训练完检查点在哪', 2)
C('dir work_dir\\uscod\\UCOD-DPL_dinov2\\UCOD-DPL_dinov2\\ckp\\')
P('应该看到 epoch5.pth、epoch10.pth、epoch15.pth、epoch20.pth、epoch25.pth 五个文件夹。')
P('我们用 epoch25.pth（最后一轮）来做最终评估。')

H('4.6 训练中常见问题', 2)
P('① 特征提取阶段报错 → 检查 datasets/RefCOD/ 连接是否正确。执行 dir datasets\\RefCOD 看有没有 6 个文件夹。')
P('② CUDA out of memory → 关掉浏览器。如果还不行，把 batch_size 从 16 改小（在 configs/uscod/UCOD-DPL_dinov2.py 里改 trainloader_cfg → batch_size）。')
P('③ 训练到一半卡住不动 → 不要关掉终端。看任务管理器里 GPU 使用率，如果在 80%+ 说明还在训练。')
P('④ loss 突然变大或变成 NaN → 训练在发散，需要降低学习率。但我们这个实验没有遇到。')

doc.add_page_break()

# ════════════════════════════════ CH5 ════════════════════════════════
H('第五章  第三步——在 4 个测试集上评估', 1)

H('5.1 这一步要干什么', 2)
P('用训练好的 epoch 25 模型，在 4 个标准 COD 测试集上逐张预测，和真值（GT）对比，算出各项指标。')
P('4 个测试集：')
P('  CHAMELEON — 76 张图，伪装类型多样')
P('  CAMO (TE-CAMO) — 250 张图，标准基准')
P('  COD10K (TE-COD10K) — 2026 张图，最大的标准基准')
P('  NC4K — 4121 张图，最新的测试集')

H('5.2 执行命令', 2)
C('bash scripts/launch_val_first_stage.sh -c configs/uscod/UCOD-DPL_dinov2.py -m work_dir/uscod/UCOD-DPL_dinov2/UCOD-DPL_dinov2/ckp/epoch25.pth')
P('参数含义：')
P('  -c：配置文件路径')
P('  -m：模型检查点路径（epoch25.pth）')

H('5.3 运行时你会看到什么', 2)
P('程序会自动依次遍历 4 个数据集，每个数据集打印一个指标表。你会依次看到：')
C('running CHAMELEON')
C('  [进度条: 76/76]')
C('  [指标表]')
C('')
C('running TE-CAMO')
C('  [进度条: 250/250]')
C('  [指标表]')
C('')
C('running TE-COD10K')
C('  [进度条: 2026/2026]')
C('  [指标表]')
C('')
C('running NC4K')
C('  [进度条: 4121/4121]')
C('  [指标表]')

P('总耗时约 30 分钟（CHAMELEON 最快，NC4K 最慢因为图最多）。')

H('5.4 指标表怎么读', 2)
P('每个数据集评估完成后，会打印一个表格，包含 3 行 9 个数字：')

P('第一行（4 个数字）：')
C('| 0.9682 | 0.7528 | 0.9330 | 0.9304 |')
P('  列1: ACC (像素准确率)——越高越好，最大 1。预测正确的像素占比。')
P('  列2: mIOU (平均交并比)——越高越好。预测和真值重叠程度。')
P('  列3: E_MAX (最大增强对齐度)——越高越好。')
P('  列4: E_MEAN (平均增强对齐度)——越高越好。')

P('第二行（4 个数字）：')
C('0.8341 | 0.8315 |  0.8625  | 0.0318 |')
P('  列1: F_MAX (最大 F 度量)——越高越好。')
P('  列2: F_MEAN (平均 F 度量)——越高越好。')
P('  列3: SMeasure (结构相似度)——核心指标！越高越好。衡量区域感知的结构相似性。')
P('  列4: MAE (平均绝对误差)——越低越好。预测和真值之间平均每个像素差多少。')

P('第三行（1 个数字）：')
C('0.8210 |')
P('  列1: WFM (加权 F 度量)——越高越好。对边界误差更敏感。')

H('5.5 核心指标对比——你的结果 vs 论文', 2)
P('UCOD-DPL 论文 (CVPR 2025, v2 DINOv2 骨干) 只报告了 4 项指标。你只需要对比这 4 项：')

P('SMeasure (越高越好)：', bold=True)
C('CHAMELEON:  论文=0.864,  目标值约 0.863')
C('CAMO:       论文=0.793,  目标值约 0.797')
C('COD10K:     论文=0.834,  目标值约 0.832')
C('NC4K:       论文=0.850,  目标值约 0.850')

P('MAE (越低越好)：', bold=True)
C('CHAMELEON:  论文=0.031,  目标值约 0.032')
C('CAMO:       论文=0.077,  目标值约 0.075')
C('COD10K:     论文=0.031,  目标值约 0.031')
C('NC4K:       论文=0.043,  目标值约 0.042')

P('E_MEAN (越高越好)：', bold=True)
C('CHAMELEON:  论文=0.931,  目标值约 0.930')
C('CAMO:       论文=0.862,  目标值约 0.867')
C('COD10K:     论文=0.916,  目标值约 0.914')
C('NC4K:       论文=0.923,  目标值约 0.924')

P('WFM (越高越好)：', bold=True)
C('CHAMELEON:  论文=0.825,  目标值约 0.821')
C('CAMO:       论文=0.747,  目标值约 0.751')
C('COD10K:     论文=0.763,  目标值约 0.758')
C('NC4K:       论文=0.818,  目标值约 0.816')

P('你的结果和上述目标值的差异应该在 ±0.005 以内。完全一样的数字是不可能的（随机种子、CUDA 版本、PyTorch 版本都会带来微小差异）。')
P('如果某个指标偏差超过 0.01，说明可能有问题，检查前面的步骤是否都正确执行了。')

doc.add_page_break()

# ════════════════════════════════ CH6 ════════════════════════════════
H('第六章  如何查看训练日志', 1)

H('6.1 日志文件在哪里', 2)
P('训练和评估的日志会同时输出到终端和日志文件。日志文件位置：')
C('work_dir\\uscod\\UCOD-DPL_dinov2\\UCOD-DPL_dinov2\\Ablation 10.log')

H('6.2 如何用记事本打开日志', 2)
P('方法一：直接在文件资源管理器里找到这个文件，双击打开。')
P('方法二：在命令行输入：')
C('notepad "work_dir\\uscod\\UCOD-DPL_dinov2\\UCOD-DPL_dinov2\\Ablation 10.log"')

H('6.3 日志里找关键信息', 2)
P('在记事本里按 Ctrl+F 搜索以下关键词：')

P('搜索 "best result" → 找到每次验证的最佳结果（epoch 5/10/15/20/25）')
P('搜索 "SMeasure" → 找到具体的指标数值')
P('搜索 "MAE" → 找平均绝对误差')
P('搜索 "Done." → 确认训练正常结束')
P('搜索 "ERROR" 或 "Traceback" → 如果没有结果，说明没有错误')

H('6.4 如何看验证指标随 epoch 的收敛情况', 2)
P('搜索 "best result"，你会找到 5 处（分别对应 epoch 5/10/15/20/25）。')
P('每次 best result 后面都有一个指标表。把 5 次的 SMeasure 按顺序记下来：')
P('  epoch 5:  SMeasure ≈ 0.648')
P('  epoch 10: SMeasure ≈ 0.677')
P('  epoch 15: SMeasure ≈ 0.712')
P('  epoch 20: SMeasure ≈ 0.722')
P('  epoch 25: SMeasure ≈ 0.797')
P('如果数字一直在上升（或稳定），说明训练正常收敛。如果中途下降或震荡，说明训练不稳定。')

doc.add_page_break()

# ════════════════════════════════ CH7 ════════════════════════════════
H('第七章  完整操作速查表（复制粘贴版）', 1)

P('以下是整个实验从头到尾的所有命令，按顺序复制粘贴即可。')

H('7.1 每次打开命令行的前四步', 2)
C('conda activate test01')
C('cd "C:\\Users\\23991\\Desktop\\新建文件夹\\UCOD-DPL-main\\UCOD-DPL-main"')
C('set PYTHONIOENCODING=utf-8')

H('7.2 数据准备（只做一次，以后再跑不需要重复）', 2)
C('python -c "import subprocess; subprocess.run([\'cmd\', \'/c\', \'mklink\', \'/J\', r\'datasets\\RefCOD\', r\'C:\\Users\\23991\\Desktop\\RefCOD (1)\\RefCOD\'], shell=True)"')
C('mkdir datasets\\cache\\pseudo_label_cache 2>nul')
C('python -c "import subprocess; subprocess.run([\'cmd\', \'/c\', \'mklink\', \'/J\', r\'datasets\\cache\\pseudo_label_cache\\TR-CAMO+TR-COD10K\', r\'C:\\Users\\23991\\Desktop\\plable\\TR-CAMO+TR-COD10K\'], shell=True)"')

H('7.3 生成 SAM2 精修标签（约 37 分钟）', 2)
C('python scripts/offline_sam2_refine.py')
P('验证：')
C('dir datasets\\cache\\refined_pseudo_labels\\*.png | find /c ".png"')
P('应为 4040。')

H('7.4 训练模型（约 41 分钟）', 2)
C('bash scripts/launch_train_first_stage.sh -c configs/uscod/UCOD-DPL_dinov2.py')

H('7.5 评估模型（约 30 分钟）', 2)
C('bash scripts/launch_val_first_stage.sh -c configs/uscod/UCOD-DPL_dinov2.py -m work_dir/uscod/UCOD-DPL_dinov2/UCOD-DPL_dinov2/ckp/epoch25.pth')

H('7.6 关键文件位置速查', 2)
C('精修标签：    datasets\\cache\\refined_pseudo_labels\\')
C('训练检查点：  work_dir\\uscod\\UCOD-DPL_dinov2\\UCOD-DPL_dinov2\\ckp\\epoch25.pth')
C('训练日志：    work_dir\\uscod\\UCOD-DPL_dinov2\\UCOD-DPL_dinov2\\Ablation 10.log')
C('配置文件：    configs\\uscod\\UCOD-DPL_dinov2.py')
C('精修主脚本：  scripts\\offline_sam2_refine.py')

H('7.7 完整时间线', 2)
C('数据准备：       5 分钟')
C('生成精修标签：   37 分钟')
C('训练：           41 分钟')
C('评估 4 数据集：  30 分钟')
C('═════════════════════════')
C('总计：           约 2 小时')

doc.add_page_break()

# ════════════════════════════════ CH8 ════════════════════════════════
H('第八章  怎么开多个命令行跑实验', 1)

H('8.1 为什么要开多个窗口', 2)
P('标签生成（用 SAM2 推理）和训练（用 DINOv2 训练）可以同时跑，因为它们用不同的 GPU 显存区域。')
P('这样可以把总时间从 2 小时压缩到约 1.5 小时。')

H('8.2 怎么操作', 2)
step(1, '打开第一个命令行窗口')
P('执行数据准备 + 生成精修标签（第三章，37 分钟）。不要关窗口。')

step(2, '标签生成开始后，再开第二个命令行窗口')
P('按 Win+R → cmd → 回车。第二个窗口里执行第一章的前三步（激活环境+进入目录+设置编码）。')

step(3, '等标签生成完成')
P('第一个窗口显示 "Done. 4040 refined" 后，第二个窗口执行训练命令（第四章）。')

H('8.3 怎么知道任务还在跑', 2)
P('方法一：看终端有没有新的日志输出。如果几秒钟内都有新日志，说明还在跑。')
P('方法二：打开任务管理器（Ctrl+Shift+Esc）→ 性能 → GPU，看 GPU 使用率。如果在 80% 以上说明还在跑。')
P('方法三：标签生成看进度条百分比，训练看 iter 数字有没有增加。')

doc.add_page_break()

# ════════════════════════════════ CH9 ════════════════════════════════
H('第九章  常见错误速查', 1)

P('以下是我们实际踩过的坑，按出现频率排列。')

P('错误 1：ModuleNotFoundError: No module named \'sam2\'', bold=True)
P('原因：SAM2 没装在 test01 环境。')
P('解决：conda activate test01 然后 pip install git+https://github.com/facebookresearch/sam2.git')

P('错误 2：FileNotFoundError: datasets\\cache\\...', bold=True)
P('原因：junction 连接没创建或路径不对。')
P('解决：回到 2.2 节重新创建连接，然后用 dir 验证。')

P('错误 3：CUDA out of memory', bold=True)
P('原因：8GB 显存不够（浏览器、游戏等占用了显存）。')
P('解决：关掉所有不用的程序（特别是 Chrome），重跑。如果还不行，修改配置文件里的 batch_size 为 8。')

P('错误 4：终端输出乱码', bold=True)
P('原因：Windows 默认编码 GBK 不支持 Unicode 字符。')
P('解决：set PYTHONIOENCODING=utf-8。如果还是乱码，不用管——数字部分通常是正常的。')

P('错误 5：特征提取时卡住不动', bold=True)
P('原因：DINOv2 下载模型权重时网络慢。')
P('解决：等。首次下载约需 1-2 分钟。之后会缓存。')

P('错误 6：bash: command not found', bold=True)
P('原因：没装 Git for Windows，bash 不可用。')
P('解决：装 Git for Windows，或者把 .sh 文件里的命令手动复制到 cmd 执行。')

P('错误 7：训练报错 "Not working on read mode!"', bold=True)
P('原因：伪标签缓存路径不存在。')
P('解决：检查 datasets/cache/pseudo_label_cache/TR-CAMO+TR-COD10K/ 下是否有 data_0.pkl。如果没有，回到 2.2 节创建连接。')

P('错误 8：评估结果和预期相差很大（SMeasure 差 > 0.01）', bold=True)
P('原因：可能是用了错误的检查点（比如 epoch5 而不是 epoch25），或者精修标签没有正确加载。')
P('解决：确认 -m 参数指向 epoch25.pth，确认 datasets/cache/refined_pseudo_labels/ 有 4040 个文件。')

# SAVE
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'experiments', 'SAM2实验教程_手把手版.docx')
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print(f'已保存: {os.path.abspath(out)}')

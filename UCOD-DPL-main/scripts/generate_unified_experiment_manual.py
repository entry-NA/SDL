"""Generate the unified AEEM v2 experiment manual as a DOCX file."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = PROJECT_ROOT / "docs" / "AEEM_V2_UNIFIED_EXPERIMENT_MANUAL.md"
DEFAULT_OUTPUT = (
    PROJECT_ROOT.parent.parent
    / "AEEM_v2_GitHub首次复现_消融_完整模型_验证_多随机种子_统一操作手册.docx"
)


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Title": (22, "1F4D78", 0, 12),
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.add_run("AEEM v2 GitHub 复现与多随机种子手册  |  ")
    add_page_number(footer)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F5F7")
    properties.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 8.5)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = False
    width_map = {
        2: [2700, 6660],
        3: [2200, 3580, 3580],
        5: [1800, 1890, 1890, 1890, 1890],
        6: [2300, 1412, 1412, 1412, 1412, 1412],
    }
    widths = width_map.get(column_count)
    if widths is None:
        base_width, remainder = divmod(9360, column_count)
        widths = [base_width + (1 if index < remainder else 0) for index in range(column_count)]

    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), "9360")
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), "120")
    table_properties.append(table_indent)
    table_layout = OxmlElement("w:tblLayout")
    table_layout.set(qn("w:type"), "fixed")
    table_properties.append(table_layout)

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for column_index, grid_column in enumerate(grid_columns):
        grid_column.set(qn("w:w"), str(widths[column_index]))

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.width = Inches(widths[column_index] / 1440)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths[column_index]))
            cell_margins = OxmlElement("w:tcMar")
            for side, margin_value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = OxmlElement(f"w:{side}")
                margin.set(qn("w:w"), str(margin_value))
                margin.set(qn("w:type"), "dxa")
                cell_margins.append(margin)
            cell_properties.append(cell_margins)
            cell.text = value
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, "Microsoft YaHei", 8.5)
                    if row_index == 0:
                        run.bold = True
    document.add_paragraph()


def render_markdown(source: Path, output: Path) -> None:
    if output.exists():
        raise FileExistsError(f"Refusing to overwrite existing DOCX: {output}")
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    index = 0
    in_code = False
    code_lines: list[str] = []
    first_heading = True
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and is_separator_row(lines[index + 1]):
            rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, rows)
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title" if first_heading else "Heading 1")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if first_heading else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run(line[2:].strip())
            first_heading = False
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            document.add_heading(line[5:].strip(), level=3)
        elif line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            run = paragraph.add_run(line[2:].rstrip("  "))
            run.italic = True
            run.font.color.rgb = RGBColor(89, 89, 89)
        elif re.match(r"^- ", line):
            document.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            document.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            document.add_paragraph(line.replace("`", "").replace("**", ""))
        index += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    render_markdown(args.source.resolve(), args.output.resolve())
    print(f"Generated: {args.output.resolve()}")


if __name__ == "__main__":
    main()

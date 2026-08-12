from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    r"C:\Users\23991\Desktop\新建文件夹\UCOD-DPL-main\UCOD-DPL-main"
    r"\experiments\AEEM_v2_qsemantic25_实验配置与手动操作教程.docx"
)

BLUE = "2E74B5"
NAVY = "203748"
DEEP_BLUE = "1F4D78"
GOLD = "C18A27"
GRAY = "505050"
MUTED = "6B7280"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF7E6"
LIGHT_GREEN = "EAF5EE"
LIGHT_RED = "FBEDEE"
LIGHT_GRAY = "F3F6F8"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, ascii_name="Calibri", east_asia="Microsoft YaHei", size=None):
    run.font.name = ascii_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)


def set_style_font(style, ascii_name="Calibri", east_asia="Microsoft YaHei", size=None):
    style.font.name = ascii_name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        style.font.size = Pt(size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="C9D2DC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, bottom_color=None, left_color=None, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    if bottom_color:
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), bottom_color)
        p_bdr.append(bottom)
    if left_color:
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), size)
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), left_color)
        p_bdr.append(left)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run = paragraph.add_run()
    begin_run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run = paragraph.add_run()
    instr_run._r.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run = paragraph.add_run()
    separate_run._r.append(separate)
    text = OxmlElement("w:t")
    text.text = "1"
    text_run = paragraph.add_run()
    text_run._r.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run = paragraph.add_run()
    end_run._r.append(end)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DEEP_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, size=size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, size=11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Code Block", 1)
    set_style_font(code, ascii_name="Consolas", east_asia="Microsoft YaHei", size=8.5)
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05

    note = doc.styles.add_style("Small Note", 1)
    set_style_font(note, size=9)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.12


def configure_header_footer(section):
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(header_table, [4200, 5160])
    repeat_table_header(header_table.rows[0])
    for cell in header_table.rows[0].cells:
        set_cell_margins(cell, top=0, bottom=20, start=0, end=0)
    left = header_table.cell(0, 0).paragraphs[0]
    right = header_table.cell(0, 1).paragraphs[0]
    left.paragraph_format.space_after = Pt(2)
    right.paragraph_format.space_after = Pt(2)
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = left.add_run("AEEM v2 实验手册")
    set_run_font(run, size=8.5)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    run = right.add_run("m4_camo_all_cod10k_qsemantic25_20260724_v1")
    set_run_font(run, ascii_name="Consolas", size=7.4)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_paragraph_border(left, bottom_color="D7DBE2", size="6")
    set_paragraph_border(right, bottom_color="D7DBE2", size="6")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    run = paragraph.add_run("实验配置与手动操作教程  |  ")
    set_run_font(run, size=8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" / ")
    set_run_font(run, size=8)
    add_field(paragraph, "NUMPAGES")


def add_title(doc, text, size=30, color=NAVY, after=8):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_centered(doc, text, size=12, color=GRAY, bold=False, italic=False, after=4):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first)
        first.bold = True
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items, level=0):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbers(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_callout(doc, title, text, fill=LIGHT_BLUE, accent=BLUE):
    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.left_indent = Inches(0.10)
    title_paragraph.paragraph_format.right_indent = Inches(0.08)
    title_paragraph.paragraph_format.space_before = Pt(4)
    title_paragraph.paragraph_format.space_after = Pt(0)
    title_paragraph.paragraph_format.keep_with_next = True
    set_paragraph_shading(title_paragraph, fill)
    set_paragraph_border(title_paragraph, left_color=accent, size="12")
    run = title_paragraph.add_run(title)
    set_run_font(run, size=10)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(accent)
    body_paragraph = doc.add_paragraph()
    body_paragraph.paragraph_format.left_indent = Inches(0.10)
    body_paragraph.paragraph_format.right_indent = Inches(0.08)
    body_paragraph.paragraph_format.space_before = Pt(0)
    body_paragraph.paragraph_format.space_after = Pt(8)
    body_paragraph.paragraph_format.keep_together = True
    set_paragraph_shading(body_paragraph, fill)
    set_paragraph_border(body_paragraph, left_color=accent, size="12")
    run = body_paragraph.add_run(text)
    set_run_font(run, size=10)
    return body_paragraph


def add_table(doc, headers, rows, widths, numeric_columns=None, font_size=9):
    numeric_columns = set(numeric_columns or [])
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(text))
        set_run_font(run, size=font_size)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index in numeric_columns else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, ascii_name="Consolas" if index in numeric_columns else "Calibri", size=font_size)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, code):
    paragraph = doc.add_paragraph(style="Code Block")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    paragraph._p.get_or_add_pPr().append(shading)
    set_paragraph_border(paragraph, left_color=BLUE, size="10")
    lines = code.strip("\n").splitlines()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, ascii_name="Consolas", east_asia="Microsoft YaHei", size=8.5)
        if index < len(lines) - 1:
            run.add_break()
    return paragraph


def add_note(doc, text):
    paragraph = doc.add_paragraph(style="Small Note")
    run = paragraph.add_run(text)
    set_run_font(run, size=9)
    return paragraph


def page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    configure_header_footer(section)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)
    kicker = add_centered(doc, "实验配置 / 操作手册", size=10.5, color=GOLD, bold=True, after=18)
    kicker.paragraph_format.keep_with_next = True
    add_title(doc, "AEEM v2 最终候选配置记录", size=28, after=7)
    add_centered(doc, "q_semantic Top-25% 标签剂量控制方案", size=15, color="2B5163", after=3)
    add_centered(doc, "基于 UCOD-DPL 的冻结 SAM2 离线伪标签精修", size=15, color="2B5163", after=26)
    add_centered(doc, "—  手动复现、训练、评估与排错指南  —", size=10.5, color=GOLD, after=76)
    add_centered(doc, "实验 ID", size=9.5, color=MUTED, after=2)
    add_centered(doc, "m4_camo_all_cod10k_qsemantic25_20260724_v1", size=10.5, color=NAVY, bold=True, after=8)
    add_centered(doc, "2026-07-25  ·  单随机种子最终候选", size=10, color=GRAY, after=4)
    add_centered(doc, "适用于 Windows PowerShell + Conda 环境 test01", size=9.5, color=MUTED, italic=True, after=10)

    page_break(doc)
    add_heading(doc, "1. 先看结论", 1)
    add_callout(
        doc,
        "当前推荐配置",
        "保留“离线零样本边界精修范式 + 自适应边缘感知增强机制（AEEM）”两项创新。"
        "AEEM 内部已从全图级 S 门控升级为语义定位校正、High/Medium/Low 提示路由、"
        "四因子候选质量、边界不确定带像素级融合、结构安全回退，并在训练源层面采用"
        "“TR-CAMO 全量 AEEM + TR-COD10K 的 q_semantic 最高 25% 使用 AEEM”。",
        fill=LIGHT_GREEN,
        accent="2F7D4A",
    )
    add_body(
        doc,
        "当前方案在你提供的《实验数据.xlsx》论文参考口径下，20 个可比指标中 18 项严格提高、"
        "2 项在显示精度下持平，四个数据集的 SMeasure 均高于论文参考值。它仍是单 seed 结果，"
        "正式论文结论需要多随机种子复验。",
    )
    add_heading(doc, "1.1 三套数字不要再混用", 2)
    add_table(
        doc,
        ["口径", "用途", "当前判断"],
        [
            ["本地基线", "判断代码与训练协议是否有增益", "四个 SMeasure 均不低于基线；COD10K 四位小数持平"],
            ["Excel 历史“完整”", "历史最高记录与旧方案参考", "当前四个 SMeasure 低于该历史行，不能称为历史最高"],
            ["Excel 论文行", "回答是否超过 UCOD-DPL 论文数据", "18/20 严格提高，2/20 持平，0/20 下降"],
        ],
        [1700, 3100, 4560],
        font_size=9,
    )
    add_callout(
        doc,
        "论文表述边界",
        "目前可以写“在当前单 seed、统一 Excel 论文口径下，四个测试集的五项主指标均不低于 UCOD-DPL”。"
        "暂时不要写“统计显著全面超越”，也不要把历史‘完整’行与论文行当成同一组基准。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "2. 你在原创新上做了哪些修改", 1)
    add_heading(doc, "2.1 两项创新名称保持不变", 2)
    add_numbers(
        doc,
        [
            "离线零样本边界精修范式：冻结 SAM2，放在 APM 之前，只改变离线伪标签输入源。",
            "自适应边缘感知增强机制（AEEM）：负责安全地处理空间偏移、提示误导、边界幻觉和结构碎片。",
        ],
    )
    add_heading(doc, "2.2 从旧版 AEEM 到 AEEM v2", 2)
    add_table(
        doc,
        ["旧版做法", "发现的问题", "当前修改"],
        [
            ["全图级 S 门控", "无法表达局部边界可靠性", "改为像素级边界置信图 Q(x)"],
            ["扫描 s_lower / s_upper / gamma", "评分不可辨识；gamma 曾使权重越界", "停止作为主优化方向"],
            ["强框 + 多正负点", "SAM2 近似复制粗标签", "宽松点提示、弱框、共识点多提示"],
            ["SAM2 可修改整图", "可能改坏核心、远端背景与拓扑", "只在边界不确定带内做残差修改"],
            ["单一多掩码总分", "语义、稳定性、边缘、安全性混在一起", "拆成 q_semantic / q_stability / q_edge / q_safety"],
            ["缺少结构校准", "COD10K 候选碎片较多", "限制连通域增长与额外结构质量"],
            ["所有训练源同等处理", "TR-COD10K 全量 AEEM 引起跨数据集权衡", "CAMO 全量，COD10K 只取 q_semantic Top-25%"],
            ["Local-SAM 作为主模块", "触发率低，贡献难以成立", "降为辅助消融，不作为当前主贡献"],
        ],
        [2400, 3000, 3960],
        font_size=8.6,
    )
    add_heading(doc, "2.3 旧版实现中已确认的逻辑问题", 2)
    add_bullets(
        doc,
        [
            "EdgeAlign 曾少除以 255，早期 97.9% 全采纳不能作为质量证据。",
            "gamma=50 时融合权重可能超过 1，粗标签权重 1-S 变为负数。",
            "s_upper=999 实际关闭 FULL 分支，不是关闭融合、全部 FULL。",
            "单个图像级 S 高，并不表示 SAM2 边界更接近真实目标。",
        ],
    )

    page_break(doc)
    add_heading(doc, "3. 当前 AEEM v2 技术方案", 1)
    add_callout(
        doc,
        "当前数据流",
        "DINOv2 粗标签与特征 → 语义定位校正 → High/Medium/Low 路由 → 冻结 SAM2 多提示候选 → "
        "四因子候选选择 → 边界不确定带像素级融合 → 结构校准/回退 → 训练源剂量控制 → APM/DBA/Look-Twice。",
    )
    add_heading(doc, "3.1 语义定位校正", 2)
    add_bullets(
        doc,
        [
            "从粗标签高置信前景和远端背景提取 DINOv2 前景/背景原型。",
            "按像素与两类原型的相似度差生成语义概率图，用于纠正粗标签空间偏移。",
            "可靠性由面积一致性、质心一致性、连通域一致性、区域 IoU、语义间隔五项平均。",
            "路由阈值：Low < 0.33；Medium 为 0.33–0.67；High >= 0.67。",
        ],
    )
    add_heading(doc, "3.2 自适应提示路由", 2)
    add_table(
        doc,
        ["路由", "提示方式", "通常候选数", "处理原则"],
        [
            ["High", "1 个高置信正点；point_only + weak_box", "6", "定位可信，窄边界带精修"],
            ["Medium", "最多 3 正点；增加 consensus_points 与安全负点", "9", "适度探索，多提示共识"],
            ["Low", "不调用 SAM2", "0", "直接回退粗标签"],
        ],
        [1200, 3900, 1500, 2760],
        numeric_columns={2},
        font_size=8.7,
    )
    add_heading(doc, "3.3 四因子候选质量", 2)
    add_table(
        doc,
        ["分量", "含义"],
        [
            ["q_semantic", "候选内部前景语义高、外部背景语义低"],
            ["q_stability", "不同提示产生的掩码是否一致"],
            ["q_edge", "候选边界是否得到真实图像梯度支持"],
            ["q_safety", "面积、质心、前景核心覆盖、背景核心排除的平均"],
        ],
        [2100, 7260],
        font_size=9,
    )
    add_body(doc, "候选总质量为四项等权平均；q_safety 必须不低于 0.25；最终最低候选质量为 0.35。")
    add_heading(doc, "3.4 边界不确定带与像素级融合", 2)
    add_bullets(
        doc,
        [
            "High 半径约为目标等效半径 5%，限制在 2–12 像素。",
            "Medium 半径约为目标等效半径 10%，限制在 4–20 像素。",
            "可靠前景核心与远端可靠背景被保护，SAM2 只能改动不确定边界区域。",
            "Q(x) = 候选质量 × 提示共识 × 语义一致性 × 边缘支持。",
        ],
    )
    add_heading(doc, "3.5 结构安全校准", 2)
    add_bullets(
        doc,
        [
            "最大有效连通域增长为 1。",
            "最大额外结构质量比例为 0.05。",
            "无可靠候选、无提示或结构风险时回退到 Soft-Coarse。",
        ],
    )

    add_heading(doc, "4. 最终标签剂量配置", 1)
    add_table(
        doc,
        ["训练源", "总数", "使用 AEEM", "使用 Soft-Coarse", "选择规则"],
        [
            ["TR-CAMO", "1000", "1000", "0", "全部使用 AEEM v2"],
            ["TR-COD10K", "3040", "760", "2280", "selected.q_semantic 最高 25%"],
            ["合计", "4040", "1760", "2280", "不读取 GT"],
        ],
        [1800, 1200, 1500, 1700, 3160],
        numeric_columns={1, 2, 3},
        font_size=8.8,
    )
    add_bullets(
        doc,
        [
            "3027 张 COD10K 样本有有效评分，13 张无有效候选并自动留在 Soft 组。",
            "760 张入选样本的 q_semantic 范围为 0.8362326473–0.9331809469。",
            "同分按文件名稳定排序；选择过程完全不读取训练 GT。",
        ],
    )
    add_callout(
        doc,
        "为什么不是全量 AEEM",
        "训练源隔离实验显示 TR-CAMO 的 AEEM 收益更稳定，而 TR-COD10K 的全量 AEEM 容易带来碎片与跨数据集权衡。"
        "当前策略把 SAM2 当成有剂量上限的离线边界专家，而不是无条件替换器。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    page_break(doc)
    add_heading(doc, "5. 最终训练配置与关键路径", 1)
    add_table(
        doc,
        ["配置项", "当前值"],
        [
            ["DINOv2", "facebook/dinov2-base"],
            ["输入尺寸", "518 × 518"],
            ["解码特征尺寸", "68"],
            ["Batch size / Epoch", "16 / 25"],
            ["学生 / 判别器学习率", "2e-4 / 1e-3"],
            ["EMA", "0.99"],
            ["APM merge", "dis"],
            ["Look-Twice", "开启；look_twice_th=0.15；expand_type=dynamic"],
            ["随机种子", "42（scripts/train.py 固定）"],
            ["混合精度", "Accelerate 启动参数 fp16"],
            ["DataLoader workers", "0"],
        ],
        [2600, 6760],
        font_size=9,
    )
    add_heading(doc, "5.1 项目目录", 2)
    add_code(doc, r'C:\Users\23991\Desktop\新建文件夹\UCOD-DPL-main\UCOD-DPL-main')
    add_heading(doc, "5.2 标签工件", 2)
    add_code(doc, r'artifacts\aeem_v2\m4_camo_all_cod10k_qsemantic25_20260724_v1')
    add_heading(doc, "5.3 最终检查点", 2)
    add_code(
        doc,
        'work_dir\\uscod\\UCOD-DPL_dinov2_aeem_v2_full4040\\'
        'UCOD-DPL_dinov2_aeem_v2_m4_camo_all_cod10k_qsemantic25_20260724_v1\\'
        r'ckp\epoch25.pth',
    )
    add_note(doc, "epoch25.pth 在当前工程中是目录式检查点；实际权重文件为该目录下的 model.safetensors。")
    add_heading(doc, "5.4 可追溯信息", 2)
    add_bullets(
        doc,
        [
            "最终混合标签输出哈希：073abc4dcd13eaa24eb050a7dc063a88dda1a5a644750eeefd8b7270cb92895e。",
            "全量 AEEM 输出哈希：2e3a081f55d806b3530d00626c231c0e221dd33ac0a0515e308e8fc6e2473850。",
            "工件记录 commit：7b7ca16e05bc34ee4fd7057541ce5f15b6ec8ae3。",
            "生成时工作区不是 clean 状态；工件内 git_diff.patch 也是复现必需证据。",
        ],
    )

    add_heading(doc, "6. 手动运行：最快复现当前配置", 1)
    add_callout(
        doc,
        "适用场景",
        "当前 m2、m0、m4 工件已经存在，只需要复用现有标签重新训练和评估。"
        "这是你现在最应该使用的路线，不需要先删除旧数据。",
        fill=LIGHT_GREEN,
        accent="2F7D4A",
    )
    add_heading(doc, "Step 0：打开 PowerShell 并进入项目目录", 2)
    add_code(
        doc,
        r'''conda activate test01
Set-Location "C:\Users\23991\Desktop\新建文件夹\UCOD-DPL-main\UCOD-DPL-main"''',
    )
    add_heading(doc, "Step 1：检查并准备最终 4040 张混合标签", 2)
    add_code(doc, r'& .\scripts\prepare_aeem_v2_qsemantic25.ps1')
    add_body(doc, "如果同名工件已经完整，脚本会显示 Already complete, skipping。这是正常行为，不是错误。")
    add_heading(doc, "Step 2：训练 25 轮", 2)
    add_code(
        doc,
        r'''& .\scripts\run_aeem_v2_train.ps1 `
  -ExperimentId m4_camo_all_cod10k_qsemantic25_20260724_v1 `
  -Port 11151''',
    )
    add_heading(doc, "Step 3：评估四个数据集", 2)
    add_code(
        doc,
        r'''& .\scripts\run_aeem_v2_eval.ps1 `
  -ExperimentId m4_camo_all_cod10k_qsemantic25_20260724_v1 `
  -Checkpoint "work_dir\uscod\UCOD-DPL_dinov2_aeem_v2_full4040\UCOD-DPL_dinov2_aeem_v2_m4_camo_all_cod10k_qsemantic25_20260724_v1\ckp\epoch25.pth" `
  -Port 11152''',
    )
    add_callout(
        doc,
        "不要清理旧实验",
        "标签工件、缓存和 checkpoint 都按 ExperimentId 隔离。直接在评估完成后的新提示行粘贴下一条命令没有问题。"
        "除非明确重建一个新实验，否则不要删除 artifacts、datasets/cache 或 work_dir。",
        fill=LIGHT_RED,
        accent="A33A3A",
    )

    page_break(doc)
    add_heading(doc, "7. 手动运行：从头重建一套新工件", 1)
    add_callout(
        doc,
        "重要",
        "全量生成脚本为保护实验不可变性，会拒绝覆盖已有目录。从头重建时必须使用新的实验 ID。"
        "下面示例用 20260725_rebuild_v1；实际运行时可替换为新的日期后缀。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )
    add_heading(doc, "Step 0：预检环境", 2)
    add_code(
        doc,
        r'''conda activate test01
Set-Location "C:\Users\23991\Desktop\新建文件夹\UCOD-DPL-main\UCOD-DPL-main"
nvidia-smi
& "C:\Anaconda\envs\test01\python.exe" -c "import torch; print(torch.cuda.is_available(), torch.cuda.get_device_name(0))"''',
    )
    add_bullets(
        doc,
        [
            "确认 datasets/RefCOD 可访问。",
            "确认 DINOv2 特征缓存和粗伪标签缓存完整。",
            "确认 SAM2.1 tiny checkpoint 路径存在。",
        ],
    )
    add_heading(doc, "Step 1：生成 Soft-Coarse 控制标签", 2)
    add_code(
        doc,
        r'''& "C:\Anaconda\envs\test01\python.exe" -u scripts\prepare_aeem_controls.py `
  --experiment-id m0_controls_20260725_rebuild_v1''',
    )
    add_heading(doc, "Step 2：生成全量 AEEM v2 标签", 2)
    add_code(
        doc,
        r'''& .\scripts\run_aeem_v2_full4040.ps1 `
  -ExperimentId m2_full4040_structure_20260725_rebuild_v1 `
  -PostprocessWorkers 2 `
  -PipelineBuffer 4''',
    )
    add_body(
        doc,
        "该步骤使用 GPU 运行 SAM2，并让 CPU 后处理与 GPU 推理通过流水线同时进行。"
        "结束时应得到 4040 张 refined_pseudo_labels、audit.jsonl、config.json、manifest.json 和哈希文件。",
    )
    add_heading(doc, "Step 3：组成 CAMO 全量 + COD10K Top-25% 标签", 2)
    add_code(
        doc,
        r'''& .\scripts\prepare_aeem_v2_qsemantic25.ps1 `
  -AeemExperimentId m2_full4040_structure_20260725_rebuild_v1 `
  -ControlExperimentId m0_controls_20260725_rebuild_v1 `
  -CamoExperimentId m4_camo_all_cod10k_qsemantic25_20260725_rebuild_v1''',
    )
    add_heading(doc, "Step 4：训练新实验", 2)
    add_code(
        doc,
        r'''& .\scripts\run_aeem_v2_train.ps1 `
  -ExperimentId m4_camo_all_cod10k_qsemantic25_20260725_rebuild_v1 `
  -Port 11161''',
    )
    add_heading(doc, "Step 5：评估新实验", 2)
    add_code(
        doc,
        r'''& .\scripts\run_aeem_v2_eval.ps1 `
  -ExperimentId m4_camo_all_cod10k_qsemantic25_20260725_rebuild_v1 `
  -Checkpoint "work_dir\uscod\UCOD-DPL_dinov2_aeem_v2_full4040\UCOD-DPL_dinov2_aeem_v2_m4_camo_all_cod10k_qsemantic25_20260725_rebuild_v1\ckp\epoch25.pth" `
  -Port 11162''',
    )

    add_heading(doc, "8. 进度条、GPU 与耗时说明", 1)
    add_table(
        doc,
        ["阶段", "主要硬件", "本机参考耗时", "正常进度表现"],
        [
            ["Soft 控制标签", "CPU + 磁盘", "较短", "文件组合/保存进度"],
            ["全量 AEEM v2", "GPU SAM2 + CPU 后处理并行", "约 2小时35分", "Refining 0/4040 → 4040/4040"],
            ["Top-25% 组合", "CPU + 磁盘", "约 30 秒", "Composing 0/4040 → 4040/4040"],
            ["25 轮训练", "GPU FP16", "约 53 分钟", "Train Epoch / Iteration"],
            ["四集评估", "GPU + CPU 数据读取", "约 27 分钟", "每个数据集一条 tqdm"],
        ],
        [1700, 2500, 1600, 3560],
        font_size=8.6,
    )
    add_heading(doc, "8.1 为什么精修时 GPU 看起来间歇工作", 2)
    add_body(
        doc,
        "SAM2 图像编码和掩码推理由 GPU 完成；图像读取、候选质量计算、结构校准、PNG 保存和审计写入由 CPU/磁盘完成。"
        "当前脚本用 2 个 CPU 后处理 worker 和 4 个流水线 buffer 重叠执行，所以 GPU 利用率可能呈波动，不代表退回纯 CPU。",
    )
    add_heading(doc, "8.2 另开窗口观察 GPU", 2)
    add_code(doc, "nvidia-smi -l 2")
    add_heading(doc, "8.3 训练结束仍显示 0% 的旧进度行", 2)
    add_body(
        doc,
        "Rich 进度组件可能在结束时保留 Train Iteration、Validation Iteration 的空行。判断是否完成应看 Train Epoch 100%、"
        "Training completed 和 checkpoint 路径，而不是最后一条残留的 0% 行。",
    )

    page_break(doc)
    add_heading(doc, "9. 每一步完成后检查什么", 1)
    add_table(
        doc,
        ["阶段", "必须检查", "通过标准"],
        [
            ["全量 AEEM", "manifest.json", "status=complete；input_count=4040；output_count=4040"],
            ["最终混合标签", "manifest.json", "source_counts：aeem=1760，soft=2280"],
            ["训练", "ckp/epoch25.pth", "目录存在，内部有 model.safetensors"],
            ["评估", "四张 Log Table", "CHAMELEON、TE-CAMO、TE-COD10K、NC4K 均输出"],
        ],
        [1900, 2700, 4760],
        font_size=8.8,
    )
    add_heading(doc, "9.1 快速检查最终标签工件", 2)
    add_code(
        doc,
        r'''$artifact = "artifacts\aeem_v2\m4_camo_all_cod10k_qsemantic25_20260724_v1"
Get-Content "$artifact\manifest.json"
(Get-ChildItem "$artifact\refined_pseudo_labels" -Filter *.png -File).Count''',
    )
    add_heading(doc, "9.2 快速检查 checkpoint", 2)
    add_code(
        doc,
        r'''$checkpoint = "work_dir\uscod\UCOD-DPL_dinov2_aeem_v2_full4040\UCOD-DPL_dinov2_aeem_v2_m4_camo_all_cod10k_qsemantic25_20260724_v1\ckp\epoch25.pth"
Test-Path $checkpoint
Get-ChildItem $checkpoint''',
    )

    add_heading(doc, "10. 常见错误与处理", 1)
    add_table(
        doc,
        ["现象", "原因", "处理"],
        [
            ["无法识别 .\\scripts\\xxx.ps1", "当前目录在项目外层", "先执行 Set-Location 到双层 UCOD-DPL-main 内层目录"],
            ["Artifact already exists", "全量脚本禁止覆盖", "复用现有工件，或换新的 ExperimentId；不要直接删旧工件"],
            ["Already complete, skipping", "目标工件已完整", "正常，继续训练即可"],
            ["端口占用", "Accelerate 主进程端口冲突", "把 -Port 改为新的未占用值"],
            ["Checkpoint not found", "路径少了一层或实验 ID 不一致", "确认 epoch25.pth 目录完整存在"],
            ["GPU 利用率为 0 一段时间", "CPU 后处理/磁盘保存阶段", "观察 1–2 分钟和 tqdm 是否继续；另开 nvidia-smi -l 2"],
            ["进度条看似停住", "单张大图或评估 Look-Twice 较慢", "先观察计数和时间；不立即中断"],
            ["评估出现 [SKIP] Logger warning", "同进程重复创建 logger 名称", "若后续权重加载和 tqdm 正常，可忽略"],
            ["PowerShell 执行策略阻止脚本", "本进程不允许执行 ps1", "Set-ExecutionPolicy -Scope Process Bypass"],
        ],
        [2500, 2900, 3960],
        font_size=8.3,
    )

    page_break(doc)
    add_heading(doc, "11. 当前评估结果与论文对比", 1)
    current_rows = [
        ["CHAMELEON", "0.9316", "0.8648", "0.0310", "0.8259", "0.8400"],
        ["TE-CAMO", "0.8639", "0.7939", "0.0760", "0.7482", "0.7811"],
        ["TE-COD10K", "0.9160", "0.8344", "0.0302", "0.7633", "0.7805"],
        ["NC4K", "0.9240", "0.8513", "0.0415", "0.8190", "0.8376"],
    ]
    add_table(
        doc,
        ["数据集", "E_MEAN", "SMeasure", "MAE", "WFM", "F_MAX"],
        current_rows,
        [2100, 1452, 1452, 1452, 1452, 1452],
        numeric_columns={1, 2, 3, 4, 5},
        font_size=8.6,
    )
    add_heading(doc, "11.1 Excel 中的 UCOD-DPL 论文参考值", 2)
    paper_rows = [
        ["CHAMELEON", "0.931", "0.864", "0.031", "0.825", "0.838"],
        ["TE-CAMO", "0.862", "0.793", "0.077", "0.747", "0.779"],
        ["TE-COD10K", "0.916", "0.834", "0.031", "0.763", "0.779"],
        ["NC4K", "0.923", "0.850", "0.043", "0.818", "0.835"],
    ]
    add_table(
        doc,
        ["数据集", "E_MEAN", "SMeasure", "MAE", "WFM", "F_MAX"],
        paper_rows,
        [2100, 1452, 1452, 1452, 1452, 1452],
        numeric_columns={1, 2, 3, 4, 5},
        font_size=8.6,
    )
    add_note(doc, r"来源：C:\Users\23991\Desktop\实验数据.xlsx，Sheet1 第 40–45 行。MAE 越低越好，其余指标越高越好。")
    add_heading(doc, "11.2 当前方案相对论文的变化", 2)
    delta_rows = [
        ["CHAMELEON", "+0.0006", "+0.0008", "持平", "+0.0009", "+0.0020"],
        ["TE-CAMO", "+0.0019", "+0.0009", "改善 0.0010", "+0.0012", "+0.0021"],
        ["TE-COD10K", "持平", "+0.0004", "改善 0.0008", "+0.0003", "+0.0015"],
        ["NC4K", "+0.0010", "+0.0013", "改善 0.0015", "+0.0010", "+0.0026"],
    ]
    add_table(
        doc,
        ["数据集", "E_MEAN", "SMeasure", "MAE", "WFM", "F_MAX"],
        delta_rows,
        [2100, 1452, 1452, 1452, 1452, 1452],
        numeric_columns={1, 2, 3, 4, 5},
        font_size=8.2,
    )
    add_callout(
        doc,
        "汇总结论",
        "20 个可比指标中：18 项严格优于论文，2 项在当前显示精度下持平，0 项下降。"
        "四个 SMeasure 分别提高 0.0008、0.0009、0.0004、0.0013。",
        fill=LIGHT_GREEN,
        accent="2F7D4A",
    )

    add_heading(doc, "12. 下一步实验顺序", 1)
    add_numbers(
        doc,
        [
            "固定当前标签与训练配置，增加至少 3 个 seed，报告均值和标准差。",
            "做 COD10K 随机 25% 对照，证明 q_semantic 排序优于随机选 760 张。",
            "做 COD10K 0%、25%、50%、100% 剂量消融，验证跨数据集权衡。",
            "同协议复跑本地基线和 Excel 历史‘完整’方案，排除 checkpoint、Look-Twice、评估脚本差异。",
            "多 seed 稳定后，再把当前候选升级为论文正式主结果。",
        ],
    )
    add_callout(
        doc,
        "建议优先级",
        "下一组最有价值的实验不是继续调 s_lower / s_upper / gamma，而是 q_semantic Top-25% 对随机 25% 的控制实验。"
        "这能直接证明最终新增的“训练源剂量控制”确实利用了语义质量，而不是仅靠减少 AEEM 使用量。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "附录 A：关键文件索引", 1)
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["scripts/run_aeem_v2_full4040.ps1", "全量 4040 张 AEEM v2 生成入口"],
            ["scripts/prepare_aeem_v2_qsemantic25.ps1", "组成最终 CAMO 全量 + COD10K Top-25% 标签"],
            ["scripts/run_aeem_v2_train.ps1", "FP16 单 GPU 训练入口"],
            ["scripts/run_aeem_v2_eval.ps1", "四个测试集统一评估入口"],
            ["aeem_v2/semantic.py", "语义定位校正与路由"],
            ["aeem_v2/refinement.py", "提示、候选评估、边界融合"],
            ["aeem_v2/structure.py", "结构校准与回退"],
            ["aeem_v2/composition.py", "按数据源与 q_semantic 组成不可变标签工件"],
            ["docs/AEEM_V2_QSEMANTIC25_FINAL_CONFIG.md", "本配置的可追溯文本记录"],
        ],
        [4000, 5360],
        font_size=8.5,
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                set_run_font(run)

    core = doc.core_properties
    core.title = "AEEM v2 q_semantic Top-25% 实验配置与手动操作教程"
    core.subject = "UCOD-DPL + 冻结 SAM2 离线伪标签精修"
    core.author = ""
    core.keywords = "AEEM v2, UCOD-DPL, SAM2, q_semantic, 手动实验"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()

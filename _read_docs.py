# -*- coding: utf-8 -*-
import docx, os, sys
sys.stdout.reconfigure(encoding='utf-8')

paths = [
    r'C:\Users\23991\Desktop\副本讨论大纲.docx',
    r'C:\Users\23991\Desktop\SAM2精修实验报告_完整版.docx',
]
for path in paths:
    print(f"\n{'='*60}")
    print(f"FILE: {os.path.basename(path)}")
    print('='*60)
    doc = docx.Document(path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            print(t)
    # Also check tables
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i+1} ---")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            print(' | '.join(cells))